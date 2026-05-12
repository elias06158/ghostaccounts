"""
GhostAccounts – Produkthandbuch Generator
Erstellt: 11. Mai 2026
Ausgabe:  Dokumente/GhostAccounts_Produkthandbuch.docx
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Farben ───────────────────────────────────────────────────────────────────
BRAND_INDIGO   = RGBColor(0x63, 0x66, 0xF1)   # #6366F1 — Indigo 500
BRAND_DARK     = RGBColor(0x1E, 0x1B, 0x4B)   # #1E1B4B — Sehr dunkles Indigo
TEXT_DARK      = RGBColor(0x11, 0x18, 0x27)   # #111827
TEXT_MUTED     = RGBColor(0x6B, 0x72, 0x80)   # #6B7280
RED            = RGBColor(0xEF, 0x44, 0x44)   # #EF4444
AMBER          = RGBColor(0xF5, 0x9E, 0x0B)   # #F59E0B
GREEN          = RGBColor(0x22, 0xC5, 0x5E)   # #22C55E
BG_LIGHT       = RGBColor(0xF8, 0xF9, 0xFF)   # Sehr heller Indigo-Ton
BORDER_INDIGO  = RGBColor(0xC7, 0xD2, 0xFE)   # Indigo 200

# ── Hilfsfunktionen ──────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex: str):
    """Setzt Hintergrundfarbe einer Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                     color="C7D2FE", size="4"):
    """Setzt Rahmen einer Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, enabled in [("top", top), ("bottom", bottom),
                          ("left", left), ("right", right)]:
        if enabled:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_break(doc):
    doc.add_page_break()


def set_run_font(run, size_pt: int, bold=False, color: RGBColor = None,
                italic=False, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text: str, level: int = 1, color: RGBColor = None,
                space_before: int = 16, space_after: int = 6):
    """Fügt eine gestylte Überschrift ein."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def add_body(doc, text: str, space_after: int = 6, color: RGBColor = None,
             bold=False, italic=False, indent: float = 0.0):
    """Fügt normalen Fließtext ein."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text: str, level: int = 0, color: RGBColor = None):
    """Fügt einen Aufzählungspunkt ein."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color
    return p


def add_info_box(doc, title: str, lines: list[str],
                 bg_hex="EEF2FF", border_hex="C7D2FE", title_color=None):
    """Erstellt eine farbige Info-Box als Tabelle."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    set_cell_borders(cell, color=border_hex)
    cell.width = Inches(6.0)

    # Titel
    if title:
        tp = cell.paragraphs[0]
        tp.paragraph_format.space_before = Pt(4)
        tp.paragraph_format.space_after  = Pt(2)
        tr = tp.add_run(title)
        tr.font.name = "Calibri"
        tr.font.size = Pt(11)
        tr.font.bold = True
        if title_color:
            tr.font.color.rgb = title_color

    for line in lines:
        lp = cell.add_paragraph()
        lp.paragraph_format.space_before = Pt(1)
        lp.paragraph_format.space_after  = Pt(1)
        lr = lp.add_run(line)
        lr.font.name = "Calibri"
        lr.font.size = Pt(10.5)
        lr.font.color.rgb = TEXT_DARK

    # Abstand nach Tabelle
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_feature_table(doc, headers: list[str], rows: list[list[str]],
                      header_bg="6366F1", header_txt=RGBColor(255,255,255)):
    """Erstellt eine gestylte Feature-Tabelle."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header-Zeile
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_bg)
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = header_txt

    # Daten-Zeilen
    for ri, row_data in enumerate(rows):
        row_cells = tbl.rows[ri + 1].cells
        bg = "F5F3FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            set_cell_bg(row_cells[ci], bg)
            p = row_cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            r.font.color.rgb = TEXT_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_risk_badge(doc, level: str, score: str, description: str):
    """Fügt ein Risiko-Badge als kleine Tabelle ein."""
    colors = {"low": "22C55E", "medium": "F59E0B", "high": "EF4444"}
    labels = {"low": "Niedriges Risiko", "medium": "Mittleres Risiko", "high": "Hohes Risiko"}
    bg = colors.get(level, "6B7280")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    c0 = tbl.cell(0, 0)
    set_cell_bg(c0, bg)
    c0.width = Inches(0.7)
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(score)
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

    c1 = tbl.cell(0, 1)
    set_cell_bg(c1, bg)
    c1.width = Inches(1.5)
    p2 = c1.paragraphs[0]
    r2 = p2.add_run(labels.get(level, ""))
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(255, 255, 255)

    c2 = tbl.cell(0, 2)
    set_cell_bg(c2, "F5F3FF")
    c2.width = Inches(3.8)
    p3 = c2.paragraphs[0]
    r3 = p3.add_run(description)
    r3.font.name = "Calibri"
    r3.font.size = Pt(10)
    r3.font.color.rgb = TEXT_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Dokumentaufbau ───────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ─── TITELSEITE ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("👻  GhostAccounts")
    r.font.name = "Calibri"
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = BRAND_INDIGO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Produkthandbuch")
    r2.font.name = "Calibri"
    r2.font.size = Pt(22)
    r2.font.color.rgb = BRAND_DARK

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Dein digitaler Datenschutz-Assistent")
    rs.font.name = "Calibri"
    rs.font.size = Pt(14)
    rs.font.italic = True
    rs.font.color.rgb = TEXT_MUTED

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        f"Version 1.0  ·  Stand: {datetime.today().strftime('%d. %B %Y')}  ·  "
        "WAMOCON GmbH  ·  Vertraulich"
    )
    rm.font.name = "Calibri"
    rm.font.size = Pt(10)
    rm.font.color.rgb = TEXT_MUTED

    add_page_break(doc)

    # ─── INHALTSVERZEICHNIS ───────────────────────────────────────────────────
    add_heading(doc, "Inhaltsverzeichnis", level=1, color=BRAND_INDIGO)
    toc_entries = [
        ("1", "Einleitung & Produktvision", "3"),
        ("2", "Zielgruppe & Anwendungsfälle", "4"),
        ("3", "Erste Schritte – Registrierung & Login", "5"),
        ("4", "Dashboard – Übersicht", "6"),
        ("5", "Scan-Funktion", "8"),
        ("5.1", "Gmail OAuth Scan", "9"),
        ("5.2", "IMAP-Scan (alle Anbieter)", "10"),
        ("5.3", "Demo-Scan", "11"),
        ("6", "Konten-Übersicht (Accounts)", "12"),
        ("6.1", "Filter & Suche", "13"),
        ("6.2", "Konto-Aktionen", "14"),
        ("6.3", "Evidenz-Ansicht", "15"),
        ("7", "Risiko-Score", "16"),
        ("8", "Risiko-Prognose (Forecast)", "18"),
        ("9", "Digital Twin – Netzwerk-Visualisierung", "19"),
        ("10", "Digital Will – Digitales Testament", "21"),
        ("11", "Einstellungen", "23"),
        ("12", "Sprachunterstützung & Lokalisierung", "25"),
        ("13", "Datenschutz & Sicherheitsarchitektur", "26"),
        ("14", "Technische Architektur", "28"),
        ("15", "Preismodelle – Free vs. Pro", "30"),
        ("16", "Glossar", "31"),
        ("17", "Rechtliche Hinweise", "32"),
    ]
    tbl = doc.add_table(rows=len(toc_entries), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (nr, title, page) in enumerate(toc_entries):
        bg = "EEF2FF" if ri % 2 == 0 else "FFFFFF"
        cells = tbl.rows[ri].cells
        for ci in range(3):
            set_cell_bg(cells[ci], bg)

        def toc_run(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=TEXT_DARK):
            p = cell.paragraphs[0]
            p.alignment = align
            r = p.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(10.5)
            r.font.bold = bold
            r.font.color.rgb = color

        toc_run(cells[0], nr, bold=bool(len(nr) == 1), color=BRAND_INDIGO)
        toc_run(cells[1], title, bold=(len(nr) == 1))
        toc_run(cells[2], page, align=WD_ALIGN_PARAGRAPH.RIGHT, color=TEXT_MUTED)

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. EINLEITUNG
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1  Einleitung & Produktvision", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "GhostAccounts ist eine datenschutzorientierte Web-Applikation, die Nutzern hilft, "
        "ihren digitalen Fußabdruck zu entdecken, zu verstehen und zu kontrollieren. Die Anwendung "
        "analysiert E-Mail-Metadaten und identifiziert Online-Dienste, bei denen ein Nutzer registriert "
        "ist – darunter auch längst vergessene Accounts, die ein Sicherheitsrisiko darstellen können.")

    add_body(doc,
        "Im Kern beantwortet GhostAccounts drei fundamentale Fragen des digitalen Zeitalters:")

    add_bullet(doc, "Bei welchen Online-Diensten bin ich registriert?")
    add_bullet(doc, "Welche davon wurden durch Datenlecks kompromittiert?")
    add_bullet(doc, "Wie kann ich meinen digitalen Nachlass regeln?")

    add_body(doc, "")

    add_info_box(doc, "🎯  Produktvision", [
        "GhostAccounts gibt Nutzern die vollständige Kontrolle über ihre digitale Identität zurück.",
        "Die App erkennt, bewertet und hilft beim Aufräumen vergessener Online-Accounts –",
        "sicher, transparent und DSGVO-konform.",
    ], bg_hex="EEF2FF", border_hex="6366F1", title_color=BRAND_INDIGO)

    add_heading(doc, "1.1  Kernprinzipien", level=2, color=BRAND_DARK)
    principles = [
        ("🔒  Privacy First",
         "GhostAccounts speichert niemals E-Mail-Inhalte. Nur abgeleitete Service-Metadaten "
         "(Domain, Erkennungstyp, Datum) werden gespeichert."),
        ("🏛️  DSGVO-Konform",
         "Alle Daten werden in der EU gespeichert (Supabase, West EU). Nutzer können ihre "
         "Daten jederzeit vollständig exportieren oder löschen."),
        ("🔍  Transparente Erkennung",
         "Jeder erkannte Account zeigt seine Evidenz (E-Mail-Typ, Datum, Erkennungsquelle). "
         "Kein 'Magic Black Box' – Nutzer sehen genau, warum ein Account erkannt wurde."),
        ("🌐  Zweisprachig",
         "Die gesamte UI ist auf Deutsch und Englisch verfügbar. Sprachumschaltung jederzeit "
         "im Header möglich."),
    ]
    for title, desc in principles:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(f"{title}:  ")
        r1.font.name = "Calibri"; r1.font.size = Pt(11); r1.font.bold = True
        r1.font.color.rgb = BRAND_INDIGO
        r2 = p.add_run(desc)
        r2.font.name = "Calibri"; r2.font.size = Pt(11)
        r2.font.color.rgb = TEXT_DARK

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. ZIELGRUPPE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2  Zielgruppe & Anwendungsfälle", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "GhostAccounts richtet sich an alle Internetnutzer, die über Jahre hinweg Online-Dienste "
        "genutzt haben und den Überblick verloren haben. Im Fokus stehen drei Hauptgruppen:")

    add_feature_table(doc,
        headers=["Zielgruppe", "Bedürfnis", "Hauptfunktion"],
        rows=[
            ["Datenschutz-bewusste Nutzer",
             "Wissen, welche Dienste persönliche Daten speichern",
             "Account-Entdeckung + Löschlinks"],
            ["Sicherheits-orientierte Nutzer",
             "Breach-Monitoring für alle Accounts",
             "HIBP-Integration + Risiko-Score"],
            ["Digitale Nachlassplanung",
             "Regelung was mit Accounts nach dem Tod passiert",
             "Digital Will"],
            ["Technik-affine Nutzer",
             "Vollständige Datenkontrolle und Export",
             "JSON/CSV-Export + Account-Löschung"],
            ["Unternehmensnutzer (Pro)",
             "Unbegrenzte Accounts + aktives Monitoring",
             "Pro-Plan Features"],
        ]
    )

    add_heading(doc, "2.1  Typische Anwendungsszenarien", level=2, color=BRAND_DARK)

    scenarios = [
        ("Szenario 1: Der vergessene Account",
         "Ein Nutzer hat sich vor 8 Jahren bei einem Online-Shop registriert, der inzwischen "
         "gehackt wurde. GhostAccounts erkennt den Account über eine alte Rechnungs-E-Mail, "
         "zeigt den Breach-Status und verlinkt direkt auf die Löschseite des Dienstes."),
        ("Szenario 2: Der digitale Frühjahrsputz",
         "Nach Jahren im Internet möchte ein Nutzer wissen, bei welchen 200+ Diensten er "
         "angemeldet ist. Der IMAP-Scan analysiert 2.000 E-Mails in unter 30 Sekunden und "
         "erstellt eine priorisierte Liste aller gefundenen Accounts."),
        ("Szenario 3: Digitales Testament",
         "Ein Nutzer möchte regeln, was mit seinen Online-Accounts passiert, falls er 6 Monate "
         "inaktiv ist. Er aktiviert den Digital Will, bestimmt dass Social-Accounts gelöscht "
         "und Bank-Accounts an seinen Erben übertragen werden sollen."),
    ]
    for title, text in scenarios:
        add_body(doc, title, bold=True, color=BRAND_DARK, space_after=2)
        add_body(doc, text, indent=0.2, space_after=8)

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. ERSTE SCHRITTE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3  Erste Schritte – Registrierung & Login", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "GhostAccounts ist als Web-Applikation unter ghostaccounts.vercel.app erreichbar. "
        "Es ist keine Installation erforderlich. Die App funktioniert auf Desktop, Tablet und "
        "mobilen Geräten vollständig responsiv.")

    add_heading(doc, "3.1  Registrierung", level=2, color=BRAND_DARK)
    add_body(doc, "Neue Nutzer können sich auf zwei Wegen registrieren:")

    add_bullet(doc, "E-Mail & Passwort:  Unter /register E-Mail-Adresse und Passwort eingeben. "
               "Eine Bestätigungs-E-Mail wird automatisch versendet.")
    add_bullet(doc, "Google OAuth:  Schnell-Registrierung per Google-Konto. "
               "Gibt zusätzlich die Möglichkeit, direkt den Gmail-Scan zu starten.")

    add_info_box(doc, "ℹ️  Sicherheitshinweis", [
        "Passwörter werden nie im Klartext gespeichert (Supabase Auth, bcrypt).",
        "Google OAuth verwendet OAuth 2.0 + PKCE – keine Passwörter werden übertragen.",
        "E-Mail-Adressen werden ausschließlich für die Authentifizierung verwendet.",
    ], bg_hex="EEF2FF", title_color=BRAND_INDIGO)

    add_heading(doc, "3.2  Login", level=2, color=BRAND_DARK)
    add_body(doc,
        "Eingeloggte Nutzer werden automatisch zum Dashboard weitergeleitet. "
        "Die Session bleibt bestehen bis zum expliziten Logout. "
        "Bei vergessenem Passwort steht eine 'Passwort vergessen'-Funktion zur Verfügung, "
        "die eine Zurücksetz-E-Mail versendet.")

    add_heading(doc, "3.3  Navigationsstruktur", level=2, color=BRAND_DARK)
    add_body(doc, "Nach dem Login stehen folgende Bereiche zur Verfügung:")

    add_feature_table(doc,
        headers=["Bereich", "URL", "Beschreibung"],
        rows=[
            ["Dashboard",       "/dashboard",          "Übersicht, Risiko-Score, Quick Actions"],
            ["Accounts",        "/dashboard/accounts", "Alle gefundenen Konten"],
            ["Scan",            "/dashboard/scan",     "E-Mail-Scan starten"],
            ["Risiko-Prognose", "/dashboard/forecast", "6-Monats-Risikoentwicklung"],
            ["Digital Twin",    "/dashboard/twin",     "Netzwerk-Visualisierung"],
            ["Digital Will",    "/dashboard/will",     "Digitales Testament"],
            ["Einstellungen",   "/dashboard/settings", "Profil, Datenschutz, Export"],
        ]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. DASHBOARD
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4  Dashboard – Übersicht", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "Das Dashboard ist die zentrale Anlaufstelle der Anwendung. Es bietet eine sofortige "
        "Übersicht über den aktuellen Sicherheitsstatus und ermöglicht schnellen Zugriff auf "
        "alle wichtigen Funktionen.")

    add_heading(doc, "4.1  Begrüßungs-Header", level=2, color=BRAND_DARK)
    add_body(doc,
        "Der Header zeigt eine personalisierte Begrüßung mit Vor- oder Benutzername sowie "
        "das Datum des letzten Scans. Wenn noch kein Scan durchgeführt wurde, wird ein "
        "prominenter Call-to-Action für den ersten Scan angezeigt.")

    add_heading(doc, "4.2  Statistik-Karten", level=2, color=BRAND_DARK)
    add_body(doc, "Nach dem ersten Scan werden vier Statistik-Karten angezeigt:")

    add_feature_table(doc,
        headers=["Karte", "Farbe", "Inhalt"],
        rows=[
            ["Risiko-Score",     "Indigo/Kreisdiagramm", "Numerischer Score 0–100 mit Kategorie (Niedrig / Mittel / Hoch)"],
            ["Accounts gefunden", "Grau",   "Gesamtzahl aller erkannten Online-Konten"],
            ["Kompromittiert",   "Rot",    "Anzahl der Accounts mit bestätigtem Datenleck"],
            ["Inaktiv (3+ Jahre)", "Amber", "Accounts ohne E-Mail-Aktivität seit über 3 Jahren"],
        ]
    )

    add_heading(doc, "4.3  Quick Actions", level=2, color=BRAND_DARK)
    add_body(doc, "Drei Schnellzugriff-Kacheln ermöglichen direkten Zugang zu den häufigsten Aktionen:")

    add_bullet(doc, "🔍  Neuen Scan starten – führt direkt zum Scan-Bereich")
    add_bullet(doc, "⚠️  Breaches prüfen – öffnet gefilterte Account-Liste (nur Breached)")
    add_bullet(doc, "✨  Aufräumen – öffnet gefilterte Account-Liste (nur Aktiv)")

    add_heading(doc, "4.4  Breach-Alerts", level=2, color=BRAND_DARK)
    add_body(doc,
        "Der Abschnitt 'Aktuelle Breach-Warnungen' zeigt die 5 neuesten ungelesenen "
        "Sicherheitswarnungen von Have I Been Pwned (HIBP). Wenn keine Warnungen vorliegen, "
        "wird eine beruhigende 'Alles sicher'-Meldung angezeigt.")

    add_heading(doc, "4.5  Ghost Risk Forecast (Vorschau)", level=2, color=BRAND_DARK)
    add_body(doc,
        "Am unteren Ende des Dashboards befindet sich eine kompakte Vorschau der Risiko-Prognose "
        "für die nächsten 6 Monate. Sie zeigt eine Sparkline-Grafik mit dem voraussichtlichen "
        "Risikoverlauf sowie empfohlene Sofortmaßnahmen.")

    add_heading(doc, "4.6  Digital Twin Teaser", level=2, color=BRAND_DARK)
    add_body(doc,
        "Eine hervorgehobene Link-Karte verweist auf den Digital Twin – die interaktive "
        "Netzwerk-Visualisierung aller gefundenen Accounts. Das Badge 'New' weist auf die "
        "neue Funktion hin.")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. SCAN-FUNKTION
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5  Scan-Funktion", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "Die Scan-Funktion ist das Herzstück von GhostAccounts. Sie analysiert E-Mail-Metadaten "
        "(Absender, Betreff, Datum) und identifiziert Online-Dienste, bei denen der Nutzer "
        "registriert ist. Dabei werden niemals E-Mail-Inhalte gespeichert oder an Server "
        "übertragen.")

    add_info_box(doc, "🔒  Datenschutz beim Scan", [
        "Was analysiert wird:  Absender-Domain, Betreff-Zeile, Datum der E-Mail",
        "Was NICHT gespeichert wird:  E-Mail-Inhalte, Anhänge, Empfänger, CC/BCC",
        "Was gespeichert wird:  Erkannter Service-Name, Domain, Erkennungstyp, Datum",
    ], bg_hex="ECFDF5", border_hex="22C55E", title_color=RGBColor(0x06, 0x6f, 0x35))

    # 5.1 Gmail OAuth
    add_heading(doc, "5.1  Gmail OAuth Scan", level=2, color=BRAND_DARK)

    add_body(doc,
        "Der Gmail OAuth Scan verbindet sich über Google's offizielle OAuth 2.0 API mit "
        "dem Gmail-Posteingang. Dieses Verfahren ist die sicherste Scan-Methode, da kein "
        "Passwort eingegeben werden muss.")

    add_body(doc, "Technischer Ablauf:", bold=True, space_after=2)
    steps = [
        "Der Nutzer klickt auf 'Mit Gmail scannen'",
        "GhostAccounts generiert einen PKCE-Code-Verifier (RFC 7636 – verhindert Interception)",
        "Der Nutzer wird zu Google weitergeleitet und erteilt die Berechtigung 'gmail.readonly'",
        "Google sendet einen Authorization Code zurück",
        "GhostAccounts tauscht den Code gegen ein Access Token aus (Code Exchange)",
        "Die Gmail API wird aufgerufen, um Metadaten der letzten Nachrichten abzurufen",
        "Parallel: Google Connected Apps werden geprüft (OAuth-Token-basiert)",
        "Ergebnisse beider Scans werden zusammengeführt und gespeichert",
    ]
    for i, step in enumerate(steps, 1):
        add_bullet(doc, f"{i}. {step}")

    add_info_box(doc, "⚠️  Voraussetzung", [
        "Google OAuth muss in den Projekt-Einstellungen konfiguriert sein (NEXT_PUBLIC_GOOGLE_CLIENT_ID).",
        "Ohne Konfiguration erscheint eine Fehlermeldung 'Gmail nicht konfiguriert'.",
        "In der aktuellen Deployment-Version ist die Demo-Alternative immer verfügbar.",
    ], bg_hex="FFFBEB", border_hex="F59E0B", title_color=RGBColor(0xB4, 0x5C, 0x09))

    # 5.2 IMAP
    add_heading(doc, "5.2  IMAP-Scan (alle E-Mail-Anbieter)", level=2, color=BRAND_DARK)

    add_body(doc,
        "Der IMAP-Scan ermöglicht das Scannen jedes beliebigen E-Mail-Postfachs über das "
        "universelle IMAP-Protokoll (Port 993, SSL/TLS). Er funktioniert mit allen gängigen "
        "deutschen und internationalen Anbietern.")

    add_body(doc, "Unterstützte E-Mail-Anbieter:", bold=True, space_after=2)
    add_feature_table(doc,
        headers=["Anbieter", "Automatische Erkennung", "Besonderheit"],
        rows=[
            ["Gmail",           "✅ Ja", "App-Passwort erforderlich (nicht Hauptpasswort)"],
            ["Outlook/Hotmail", "✅ Ja", "App-Passwort in Microsoft-Konto generieren"],
            ["Yahoo Mail",      "✅ Ja", "App-Passwort unter security.yahoo.com"],
            ["iCloud Mail",     "✅ Ja", "App-Passwort bei appleid.apple.com"],
            ["GMX",             "✅ Ja", "Normales GMX-Passwort verwendbar"],
            ["WEB.DE",          "✅ Ja", "Normales WEB.DE-Passwort verwendbar"],
            ["T-Online",        "✅ Ja", "Normales T-Online-Passwort verwendbar"],
            ["Andere (IMAP)",   "✅ Manuell", "IMAP-Host wird automatisch als imap.[domain] gesetzt"],
        ]
    )

    add_body(doc, "Scan-Umfang und Performance:", bold=True, space_after=2)
    add_bullet(doc, "Maximal 2.000 neueste E-Mails werden analysiert")
    add_bullet(doc, "Verbindungs-Timeout: 15 Sekunden")
    add_bullet(doc, "Typische Scan-Dauer: 5–30 Sekunden je nach Postfach-Größe")
    add_bullet(doc, "SSL/TLS-verschlüsselte Verbindung (Port 993) immer aktiv")

    add_body(doc, "")
    add_info_box(doc, "🔑  App-Passwörter", [
        "Bei Anbietern mit Zwei-Faktor-Authentifizierung (Gmail, Outlook, Yahoo, iCloud)",
        "muss ein App-Passwort verwendet werden – nicht das Hauptpasswort.",
        "App-Passwörter können jederzeit wieder deaktiviert werden.",
        "GhostAccounts speichert das Passwort NICHT – es wird nur für die Scan-Dauer gehalten.",
    ], bg_hex="EEF2FF", border_hex="6366F1", title_color=BRAND_INDIGO)

    # 5.3 Demo
    add_heading(doc, "5.3  Demo-Scan", level=2, color=BRAND_DARK)

    add_body(doc,
        "Der Demo-Scan simuliert einen echten Scan mit einem vordefinierten Datensatz von "
        "30 bekannten Online-Diensten (Google, Facebook, Instagram, Amazon, etc.). "
        "Er dient zur Demonstration der Funktionen ohne eigene E-Mail-Verbindung.")

    add_body(doc,
        "Demo-Ergebnisse werden mit 'detection_source: demo' markiert und können jederzeit "
        "durch einen echten Scan überschrieben werden. Der Demo-Scan ist für alle Nutzer "
        "kostenlos verfügbar.")

    # 5.4 Erkennungslogik
    add_heading(doc, "5.4  Erkennungslogik & Evidenztypen", level=2, color=BRAND_DARK)

    add_body(doc,
        "GhostAccounts verwendet ein Gewichtungssystem zur Erkennung von Accounts. "
        "Nur E-Mails mit starken Signalen werden als bewiesene Accounts gewertet. "
        "Marketing-E-Mails und Newsletter werden explizit ignoriert, um Falsch-Positive "
        "zu minimieren.")

    add_feature_table(doc,
        headers=["Evidenztyp", "Gewichtung", "Beispiele"],
        rows=[
            ["registration",     "6 (Hoch)", "E-Mail bestätigen, Konto erstellt, Registrierung erfolgreich"],
            ["security_alert",   "6 (Hoch)", "Passwort zurücksetzen, Login-Code, 2FA-Code, neue Anmeldung"],
            ["invoice",          "5 (Hoch)", "Rechnung, Bestellbestätigung, Zahlung erfolgreich"],
            ["password_reset",   "5 (Hoch)", "Passwort-Reset-Links"],
            ["newsletter",       "0 (Ignoriert)", "Werbemails, Marketing, Gutscheine – KEIN Beweis"],
        ]
    )

    add_body(doc, "Erkennungs-Konfidenz:", bold=True, space_after=2)
    add_bullet(doc, "high:  2+ starke Signale oder 1 Registrierungs-Signal")
    add_bullet(doc, "medium:  1 starkes Signal (Rechnung, Sicherheitswarnung)")
    add_bullet(doc, "low:  Nur Absender-Domain erkannt, kein starkes Signal im Betreff")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 6. ACCOUNTS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6  Konten-Übersicht (Accounts)", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "Die Konten-Übersicht ist die zentrale Verwaltungsseite für alle erkannten Online-Accounts. "
        "Sie bietet eine vollständige Liste mit Filter-, Such- und Verwaltungsfunktionen.")

    add_heading(doc, "6.1  Statistik-Banner", level=2, color=BRAND_DARK)
    add_body(doc, "Oberhalb der Liste werden vier Kennzahlen angezeigt:")
    add_bullet(doc, "Gesamt:  Alle verifizierten Accounts (nur mit starken Signalen)")
    add_bullet(doc, "Kompromittiert:  Accounts mit Breach-Status 'breached' (rot)")
    add_bullet(doc, "Inaktiv:  Aktive Accounts ohne E-Mail-Aktivität seit 3+ Jahren (amber)")
    add_bullet(doc, "Sicher:  Aktive, nicht-kompromittierte Accounts (grün)")

    add_heading(doc, "6.2  Filter & Suche", level=2, color=BRAND_DARK)
    add_body(doc, "Status-Filter (Tabs):")
    add_feature_table(doc,
        headers=["Filter", "Zeigt"],
        rows=[
            ["Alle",             "Alle verifizierten Accounts"],
            ["Aktiv",            "Accounts mit deletion_status = active"],
            ["Kompromittiert",   "Accounts mit breach_status = breached"],
            ["Ignoriert",        "Manuell als ignoriert markierte Accounts"],
            ["Gelöscht",         "Als gelöscht markierte Accounts"],
        ]
    )

    add_body(doc, "Kategorie-Filter (Chips):")
    add_bullet(doc, "Automatische Erkennung der Kategorie anhand der Service-Domain")
    add_bullet(doc, "Kategorien: Social, Shopping, Finance, Work, Streaming, Gaming, Travel, Tools, AI, Other")
    add_bullet(doc, "Chips werden nur für Kategorien angezeigt, die in der Nutzerliste vorhanden sind")

    add_body(doc, "Suchfunktion:")
    add_bullet(doc, "Live-Suche nach Service-Name und Domain")
    add_bullet(doc, "Kombinierbar mit Status- und Kategorie-Filtern")

    add_heading(doc, "6.3  Account-Karten", level=2, color=BRAND_DARK)
    add_body(doc, "Jeder Account wird als Karte dargestellt mit folgenden Elementen:")
    add_bullet(doc, "Service-Name & Domain mit Favicon")
    add_bullet(doc, "Kategorie-Badge und Erkennungs-Konfidenz (high / medium / low)")
    add_bullet(doc, "Breach-Status (Safe / Breached / Unknown)")
    add_bullet(doc, "Letztes E-Mail-Datum und Erkennungsquelle (IMAP / Gmail / Demo)")
    add_bullet(doc, "Löschschwierigkeit (Einfach / Mittel / Schwer) wenn bekannt")

    add_heading(doc, "6.4  Konto-Aktionen", level=2, color=BRAND_DARK)
    add_feature_table(doc,
        headers=["Aktion", "Funktion"],
        rows=[
            ["Account löschen (Link)",  "Öffnet die offizielle Lösch-URL des Dienstes in neuem Tab"],
            ["Als gelöscht markieren",  "Setzt deletion_status = deleted (behält Daten lokal)"],
            ["Ignorieren",              "Setzt deletion_status = ignored (ausblendet aus Aktiv-Ansicht)"],
            ["Reaktivieren",            "Setzt deletion_status = active"],
        ]
    )

    add_heading(doc, "6.5  Evidenz-Ansicht", level=2, color=BRAND_DARK)
    add_body(doc,
        "Jeder Account kann per Klick auf 'Evidenz anzeigen' erweitert werden. "
        "Die Evidenz-Ansicht zeigt transparent, auf Basis welcher E-Mail-Signale der "
        "Account erkannt wurde:")
    add_bullet(doc, "Erkennungsquelle (IMAP / Gmail / Demo / Gemischt)")
    add_bullet(doc, "Anzahl der gefundenen E-Mails (evidence_count)")
    add_bullet(doc, "Evidenztypen (Registrierung / Rechnung / Sicherheitswarnung / Passwort-Reset)")
    add_bullet(doc, "Absender-Domains der gefundenen E-Mails")

    add_heading(doc, "6.6  Löschanleitungen", level=2, color=BRAND_DARK)
    add_body(doc,
        "Für viele bekannte Dienste sind direkte Löschlinks hinterlegt. Für Dienste "
        "ohne direkten Link erscheint eine manuelle Löschanleitung. "
        "Die Anleitung beschreibt Schritt für Schritt, wie das Konto gelöscht werden kann, "
        "inkl. Hinweis auf die Schwierigkeit (easy / medium / hard).")

    add_info_box(doc, "💡  Free vs. Pro", [
        "Free-Nutzer: Sehen maximal 20 Accounts in der Liste.",
        "Pro-Nutzer: Unlimitierte Accounts + alle Filter-Funktionen.",
        "Das Limit gilt nur für die Anzeige – erkannt werden alle Accounts.",
    ], bg_hex="EEF2FF", border_hex="6366F1", title_color=BRAND_INDIGO)

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 7. RISIKO-SCORE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7  Risiko-Score", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "Der Risiko-Score ist ein numerischer Wert von 0 bis 100, der den aktuellen "
        "Sicherheitsstatus des digitalen Fußabdrucks bewertet. Er wird bei jedem "
        "Seitenaufruf in Echtzeit aus den aktuellen Scan-Daten berechnet.")

    add_heading(doc, "7.1  Score-Kategorien", level=2, color=BRAND_DARK)

    add_risk_badge(doc, "low",    "0–34",  "Niedriges Risiko – Account-Landschaft gut unter Kontrolle")
    add_risk_badge(doc, "medium", "35–64", "Mittleres Risiko – Handlungsbedarf vorhanden")
    add_risk_badge(doc, "high",   "65–100","Hohes Risiko – sofortige Maßnahmen empfohlen")

    add_heading(doc, "7.2  Score-Berechnung", level=2, color=BRAND_DARK)
    add_body(doc, "Der Score setzt sich aus drei gewichteten Komponenten zusammen:")

    add_feature_table(doc,
        headers=["Komponente", "Max. Punkte", "Berechnung"],
        rows=[
            ["Breach-Exposition",    "35 Punkte",
             "Anteil kompromittierter Accounts × 35. Jeder Breach-Account erhöht den Score."],
            ["Inaktive Accounts",    "40 Punkte",
             "Anteil inaktiver Accounts (keine Mail >3 Jahre) × 40. Viele Karteileichen = hohes Risiko."],
            ["Account-Fußabdruck",   "25 Punkte",
             "Logarithmische Skala der Gesamtanzahl aktiver Accounts (ab 20 Accounts voller Beitrag)."],
        ]
    )

    add_body(doc, "Formel:", bold=True, space_after=2)
    formula_p = doc.add_paragraph()
    formula_p.paragraph_format.left_indent = Inches(0.3)
    formula_p.paragraph_format.space_after = Pt(8)
    rf = formula_p.add_run(
        "Score = (breachedAccounts / totalAccounts × 35) + "
        "(inactiveAccounts / totalAccounts × 40) + "
        "(min(log₁₀(totalActive + 1) / log₁₀(21), 1.0) × 25)"
    )
    rf.font.name = "Courier New"
    rf.font.size = Pt(10)
    rf.font.color.rgb = BRAND_INDIGO

    add_heading(doc, "7.3  Score-Insights", level=2, color=BRAND_DARK)
    add_body(doc,
        "Unterhalb des Scores werden 'Main Risk Drivers' angezeigt – "
        "die drei Hauptkomponenten mit ihrem aktuellen Beitrag zum Gesamtscore. "
        "Jede Komponente zeigt ihren Anteil in Punkten und gibt kontextuelle "
        "Handlungsempfehlungen.")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 8. RISIKO-PROGNOSE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8  Risiko-Prognose (Forecast)", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "Die Risiko-Prognose berechnet, wie sich der Risiko-Score in den nächsten 6 Monaten "
        "entwickeln wird, wenn der Nutzer nichts unternimmt. Die Prognose basiert auf dem "
        "aktuellen Account-Zustand und bekannten Breach-Trends.")

    add_heading(doc, "8.1  Prognose-Darstellung", level=2, color=BRAND_DARK)
    add_body(doc, "Die Seite zeigt:")
    add_bullet(doc, "Sparkline-Grafik:  SVG-basierter Linien-Graph für die Monate 0–6")
    add_bullet(doc, "Trend-Badge:  'Stabil', '+X Punkte' oder '−X Punkte'")
    add_bullet(doc, "Headline-Text:  Kurzbewertung der Prognose ('Ihr Risiko ist stabil...')")
    add_bullet(doc, "Empfohlene Aktionen:  3 priorisierte Maßnahmen zur Risikoreduktion")

    add_heading(doc, "8.2  Prognose-Algorithmus", level=2, color=BRAND_DARK)
    add_body(doc,
        "Die Prognose simuliert für jeden Monat den zukünftigen Score unter der Annahme, "
        "dass keine Maßnahmen ergriffen werden:")
    add_bullet(doc, "Accounts, die kurz vor dem 3-Jahres-Inaktivitäts-Schwellwert stehen, erhöhen den Score")
    add_bullet(doc, "Bekannte Breach-Trends (historische Leak-Frequenz) fließen ein")
    add_bullet(doc, "Die Kurvenform wird farbig kodiert: Grün (sicher) → Amber (mittel) → Rot (hoch)")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 9. DIGITAL TWIN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9  Digital Twin – Netzwerk-Visualisierung", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "Der Digital Twin ist eine interaktive Netzwerk-Visualisierung aller erkannten "
        "Online-Accounts. Er stellt den gesamten digitalen Fußabdruck als Graphen dar "
        "und macht Zusammenhänge und Risiken intuitiv sichtbar.")

    add_heading(doc, "9.1  Technische Umsetzung", level=2, color=BRAND_DARK)
    add_body(doc,
        "Die Visualisierung verwendet D3.js mit einem Force-Directed Graph Layout. "
        "Physik-basierte Simulation sorgt für eine organische, übersichtliche Anordnung "
        "der Knoten.")

    add_feature_table(doc,
        headers=["Element", "Beschreibung"],
        rows=[
            ["Zentraler Knoten",  "E-Mail-Adresse des Nutzers – blau, groß, fixiert in der Mitte"],
            ["Service-Knoten",    "Jeder erkannte Account als Kreis, Größe = Anzahl der Evidence-E-Mails"],
            ["Verbindungslinien", "Linien zeigen Verknüpfung; Farbe = Risikostufe des Accounts"],
            ["Farb-Kodierung",    "Rot = Breached, Amber = Inaktiv, Grün = Aktiv & Sicher"],
            ["Kategorie-Farben",  "Social=Pink, Shopping=Orange, Finance=Grün, Streaming=Lila, AI=Violett"],
            ["Glow-Effekt",       "Kompromittierte Accounts leuchten rot – sofort erkennbar"],
        ]
    )

    add_heading(doc, "9.2  Interaktionsmöglichkeiten", level=2, color=BRAND_DARK)
    add_bullet(doc, "Zoom:  Mausrad oder Pinch-Geste zum Vergrößern/Verkleinern (0.3× bis 3×)")
    add_bullet(doc, "Drag:  Einzelne Knoten per Drag & Drop verschieben")
    add_bullet(doc, "Pan:  Gesamten Graph verschieben")
    add_bullet(doc, "Hover/Tooltip:  Mouseover zeigt Service-Name, Domain und Risk-Level")
    add_bullet(doc, "Klick:  Öffnet die Account-Detail-Seite des jeweiligen Dienstes")
    add_bullet(doc, "Kategorie-Filter:  Chips oben filtern den Graph nach Kategorie")
    add_bullet(doc, "Legende:  Farb-Legende zum Verständnis der Risiko-Kodierung")

    add_heading(doc, "9.3  Legende & Risiko-Indikation", level=2, color=BRAND_DARK)
    add_body(doc, "Die Legende zeigt die Farbzuordnung für Risikoniveaus:")
    add_bullet(doc, "🟢  Grün:  Aktiver, nicht-kompromittierter Account")
    add_bullet(doc, "🟡  Amber:  Inaktiver Account (keine Aktivität seit 3+ Jahren)")
    add_bullet(doc, "🔴  Rot:  Kompromittierter Account (Datenleck bekannt)")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 10. DIGITAL WILL
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10  Digital Will – Digitales Testament", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "Das Digital Will (Digitales Testament) ist eine einzigartige Funktion, mit der "
        "Nutzer festlegen können, was mit ihren Online-Accounts passiert, wenn sie "
        "für eine definierte Zeit inaktiv waren. Die Funktion adressiert ein wachsendes "
        "gesellschaftliches Problem: den digitalen Nachlass.")

    add_heading(doc, "10.1  Aktivierung und Konfiguration", level=2, color=BRAND_DARK)
    add_body(doc, "Das Will wird über einen Aktivierungs-Toggle ein- und ausgeschaltet. "
             "Im aktivierten Zustand können folgende Parameter konfiguriert werden:")

    add_feature_table(doc,
        headers=["Parameter", "Optionen", "Beschreibung"],
        rows=[
            ["Inaktivitätsfrist", "1, 3, 6, 12, 24 Monate",
             "Zeitraum ohne Login, nach dem das Will ausgeführt wird"],
            ["Globale Aktion",    "Löschen / Übertragen / Archivieren",
             "Standard-Aktion für alle Accounts (kann pro Account überschrieben werden)"],
            ["Erben-E-Mail",      "E-Mail-Adresse",
             "Bei Aktion 'Übertragen': Empfänger der Zugangsdaten"],
            ["Erben-Name",        "Text",
             "Anzeigename des Erben für die Übergabe-Kommunikation"],
            ["Persönliche Nachricht", "Freitext",
             "Optionale persönliche Botschaft an den Erben oder Hinterbliebene"],
        ]
    )

    add_heading(doc, "10.2  Globale Aktionen", level=2, color=BRAND_DARK)

    actions_detail = [
        ("🗑️  Löschen",
         "Alle betroffenen Accounts werden über ihre offiziellen Lösch-URLs entfernt. "
         "GhostAccounts führt keine automatische Löschung durch – die Aktion ist als "
         "Aufgabe für den Erben oder einen Verwalter dokumentiert."),
        ("👥  Übertragen",
         "Zugangsdaten und Kontoinformationen werden an die hinterlegte Erben-E-Mail-Adresse "
         "weitergeleitet. Die persönliche Nachricht wird beigefügt."),
        ("📦  Archivieren",
         "Accounts werden in einen Inaktivitätszustand versetzt. Kein Löschen – der digitale "
         "Fußabdruck bleibt erhalten, wird aber als 'archiviert' markiert."),
    ]
    for title, desc in actions_detail:
        add_body(doc, title, bold=True, color=BRAND_DARK, space_after=2)
        add_body(doc, desc, indent=0.2, space_after=6)

    add_heading(doc, "10.3  Per-Account-Überschreibungen", level=2, color=BRAND_DARK)
    add_body(doc,
        "Für jeden einzelnen erkannten Account kann die globale Aktion individuell "
        "überschrieben werden. Dies ermöglicht differenzierte Regelungen, z.B.:")
    add_bullet(doc, "Finance-Accounts → 'Übertragen' (an Erben für steuerliche Abwicklung)")
    add_bullet(doc, "Social-Accounts → 'Löschen' (Datenschutz)")
    add_bullet(doc, "Kreativ-Accounts (GitHub, Behance) → 'Archivieren' (Werk erhalten)")

    add_body(doc,
        "Die erweiterte Account-Liste zeigt alle Accounts mit ihrem aktuellen Will-Status "
        "und ermöglicht per Dropdown-Aktion die sofortige Änderung.")

    add_heading(doc, "10.4  Statistik-Übersicht", level=2, color=BRAND_DARK)
    add_body(doc, "Eine kompakte Statistik zeigt die Verteilung der Will-Aktionen auf einen Blick:")
    add_bullet(doc, "Anzahl 'Löschen'-Accounts")
    add_bullet(doc, "Anzahl 'Übertragen'-Accounts")
    add_bullet(doc, "Anzahl 'Archivieren'-Accounts")
    add_bullet(doc, "Anzahl 'Behalten'-Accounts")
    add_bullet(doc, "Anzahl 'Erbt Global'-Accounts (noch nicht individuell konfiguriert)")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 11. EINSTELLUNGEN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11  Einstellungen", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "Der Einstellungsbereich ermöglicht die vollständige Verwaltung des Profils, "
        "der Benachrichtigungen, des App-Themes und der eigenen Daten.")

    add_heading(doc, "11.1  Profil-Verwaltung", level=2, color=BRAND_DARK)
    add_feature_table(doc,
        headers=["Feld", "Bearbeitbar", "Beschreibung"],
        rows=[
            ["Vollständiger Name",  "✅ Ja",  "Anzeigename (wird im Dashboard-Greeting verwendet)"],
            ["E-Mail-Adresse",      "❌ Nein", "Login-E-Mail – Änderung über Supabase Auth"],
            ["Plan",                "❌ Nein", "free / pro (Upgrade über Vercel/Support)"],
        ]
    )

    add_heading(doc, "11.2  Benachrichtigungen", level=2, color=BRAND_DARK)
    add_body(doc, "Zwei Benachrichtigungs-Toggles:")
    add_bullet(doc, "Breach-Warnungen:  Benachrichtigung bei neuen Datenlecks (Standard: Aktiviert)")
    add_bullet(doc,
               "Neue Accounts:  Benachrichtigung wenn neue Accounts per Scan entdeckt wurden "
               "(Standard: Deaktiviert)")

    add_heading(doc, "11.3  Erscheinungsbild (Theme)", level=2, color=BRAND_DARK)
    add_body(doc, "GhostAccounts unterstützt drei Theme-Modi:")
    add_bullet(doc, "☀️  Hell:  Weißes Theme mit dunkelgrauem Text")
    add_bullet(doc, "🌙  Dunkel:  Dunkles Theme mit hellem Text (Standard)")
    add_bullet(doc, "💻  System:  Automatische Anpassung an Betriebssystem-Einstellung")

    add_heading(doc, "11.4  Daten-Export", level=2, color=BRAND_DARK)
    add_body(doc, "Zwei Export-Formate stehen zur Verfügung:")
    add_bullet(doc,
               "JSON-Export:  Vollständiger Export aller Scan-Ergebnisse und Profildaten "
               "als strukturierte JSON-Datei (ghostaccounts-data.json)")
    add_bullet(doc,
               "CSV-Export:  Tabellarkischer Export der Account-Daten mit Feldern: "
               "service_name, service_domain, breach_status, deletion_status, "
               "detection_confidence, last_email_date, first_detected_at, evidence_count "
               "(ghostaccounts-accounts.csv)")

    add_heading(doc, "11.5  Konto-Löschung (Account-Deletion)", level=2, color=BRAND_DARK)
    add_body(doc,
        "Nutzer können ihren GhostAccounts-Account und alle zugehörigen Daten permanent löschen. "
        "Zur Sicherheit muss das Wort 'DELETE' (Englisch) oder 'LÖSCHEN' (Deutsch) in ein "
        "Bestätigungsfeld eingegeben werden. Nach Bestätigung werden gelöscht:")
    add_bullet(doc, "Alle scan_results des Nutzers")
    add_bullet(doc, "Alle breach_alerts des Nutzers")
    add_bullet(doc, "Das Profil-Eintrag in der profiles-Tabelle")
    add_bullet(doc, "Der Auth-Account (Login nicht mehr möglich)")

    add_info_box(doc, "⚠️  Unwiderruflich", [
        "Die Konto-Löschung ist NICHT rückgängig zu machen.",
        "Alle Scan-Daten, der Risiko-Score und der Digital Will werden permanent gelöscht.",
        "Ein Daten-Export vor der Löschung wird empfohlen.",
    ], bg_hex="FEF2F2", border_hex="EF4444", title_color=RED)

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 12. SPRACHUNTERSTÜTZUNG
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12  Sprachunterstützung & Lokalisierung", level=1, color=BRAND_INDIGO)

    add_body(doc,
        "GhostAccounts ist vollständig zweisprachig implementiert. Beide Sprachen sind "
        "vollständig gleichwertig – keine Funktionen sind in einer Sprache eingeschränkt.")

    add_feature_table(doc,
        headers=["Sprache", "URL-Prefix", "Lokalisierungsdatei"],
        rows=[
            ["Deutsch (DE)", "/de/...", "src/messages/de.json"],
            ["Englisch (EN)", "/en/...", "src/messages/en.json"],
        ]
    )

    add_body(doc,
        "Die Sprachumschaltung ist jederzeit über den Header-Button verfügbar. "
        "Die gewählte Sprache wird in der URL kodiert und kann direkt verlinkt werden. "
        "Datumsformate werden sprachabhängig angepasst (DD.MM.YYYY für DE, DD/MM/YYYY für EN).")

    add_body(doc,
        "Technisch basiert die Lokalisierung auf next-intl, optimiert für "
        "den Next.js App Router. Alle Übersetzungen werden serverseitig gerendert "
        "(Server Components) oder clientseitig via useTranslations()-Hook geladen.")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 13. DATENSCHUTZ & SICHERHEIT
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13  Datenschutz & Sicherheitsarchitektur", level=1, color=BRAND_INDIGO)

    add_heading(doc, "13.1  Datenminimierung", level=2, color=BRAND_DARK)
    add_body(doc, "GhostAccounts speichert nur das absolut Notwendige:")

    add_feature_table(doc,
        headers=["Datenkategorie", "Gespeichert", "Nicht gespeichert"],
        rows=[
            ["E-Mail-Inhalte",        "❌ Nein",  "Textkörper, Anhänge, HTML"],
            ["E-Mail-Metadaten",      "Teilweise", "Nur Domain, Erkennungstyp, Datum"],
            ["IMAP-Passwort",         "❌ Nein",  "Wird nur für Scan-Dauer im RAM gehalten"],
            ["OAuth-Token",           "❌ Nein",  "Kein persistentes Token-Storage"],
            ["Nutzer-E-Mail",         "✅ Ja",    "Für Authentifizierung"],
            ["Erkannte Services",     "✅ Ja",    "Service-Name, Domain, Risiko-Status"],
            ["Scan-Zeitpunkte",       "✅ Ja",    "last_scan_at im Profil"],
        ]
    )

    add_heading(doc, "13.2  Authentifizierungssicherheit", level=2, color=BRAND_DARK)
    add_bullet(doc, "Supabase Auth mit JWT-basierter Session-Verwaltung")
    add_bullet(doc, "Passwörter: bcrypt-gehashed, niemals Klartext")
    add_bullet(doc, "OAuth 2.0 + PKCE (Proof Key for Code Exchange) – verhindert Token-Interception")
    add_bullet(doc, "HTTP-Only Cookies für Session-Token (kein JavaScript-Zugriff)")
    add_bullet(doc, "Row Level Security (RLS) auf allen Datenbanktabellen")

    add_heading(doc, "13.3  Row Level Security (RLS)", level=2, color=BRAND_DARK)
    add_body(doc,
        "Alle Datenbanktabellen sind mit Supabase Row Level Security geschützt. "
        "Jede Datenbankabfrage wird auf Datenbankebene auf den authentifizierten Nutzer "
        "beschränkt. Auch mit dem öffentlichen API-Key (anon key) kann kein Nutzer "
        "die Daten eines anderen Nutzers lesen oder schreiben.")

    add_feature_table(doc,
        headers=["Tabelle", "RLS-Policy"],
        rows=[
            ["profiles",            "auth.uid() = id – SELECT, UPDATE, DELETE nur eigenes Profil"],
            ["scan_results",        "auth.uid() = user_id – alle Operationen nur eigene Daten"],
            ["breach_alerts",       "auth.uid() = user_id – alle Operationen nur eigene Daten"],
            ["digital_will",        "auth.uid() = user_id – alle Operationen nur eigenes Will"],
            ["digital_will_items",  "auth.uid() = user_id – alle Operationen nur eigene Items"],
        ]
    )

    add_heading(doc, "13.4  DSGVO-Konformität", level=2, color=BRAND_DARK)
    add_bullet(doc, "Datenspeicherung in der EU (West EU Ireland / Frankfurt)")
    add_bullet(doc, "Vollständiger Daten-Export (Recht auf Datenportabilität – Art. 20 DSGVO)")
    add_bullet(doc, "Vollständige Konto-Löschung (Recht auf Vergessenwerden – Art. 17 DSGVO)")
    add_bullet(doc, "Cookie-Banner mit Zustimmung vor Tracking")
    add_bullet(doc, "Datenschutzerklärung, Impressum und AGB verlinkt im Footer")
    add_bullet(doc, "Betreiber: WAMOCON GmbH, HRB 123666, Eschborn, Deutschland")

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 14. TECHNISCHE ARCHITEKTUR
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "14  Technische Architektur", level=1, color=BRAND_INDIGO)

    add_heading(doc, "14.1  Tech-Stack", level=2, color=BRAND_DARK)
    add_feature_table(doc,
        headers=["Schicht", "Technologie", "Version / Details"],
        rows=[
            ["Frontend Framework",   "Next.js (App Router)",     "16.x – Server Components, Streaming"],
            ["Sprache",              "TypeScript",               "Strict Mode"],
            ["Styling",              "Tailwind CSS v4",          "Utility-First, Dark Mode"],
            ["Datenbank",            "Supabase (PostgreSQL)",    "Row Level Security, Auth, Real-time"],
            ["Auth",                 "Supabase Auth",            "OAuth 2.0, PKCE, JWT, Magic Links"],
            ["Deployment",           "Vercel",                   "Edge Network, CI/CD via GitHub Actions"],
            ["i18n",                 "next-intl",                "App Router Optimiert, DE + EN"],
            ["Visualisierung",       "D3.js",                    "Force-Directed Graph"],
            ["IMAP-Client",          "imapflow",                 "Node.js IMAP Library"],
        ]
    )

    add_heading(doc, "14.2  Datenbankschema", level=2, color=BRAND_DARK)
    add_body(doc, "Die Datenbank besteht aus 5 Tabellen im public-Schema:")

    add_feature_table(doc,
        headers=["Tabelle", "Primärschlüssel", "Wichtige Felder"],
        rows=[
            ["profiles",
             "id (uuid, FK auth.users)",
             "email, full_name, plan, language, last_scan_at, notify_breach"],
            ["scan_results",
             "id (uuid)",
             "user_id, service_name, service_domain, breach_status, deletion_status, "
             "detection_source, evidence_types, evidence_count"],
            ["breach_alerts",
             "id (uuid)",
             "user_id, service_name, breach_name, data_types, is_read"],
            ["digital_will",
             "id (uuid), UNIQUE user_id",
             "is_active, inactivity_months, global_action, transfer_email"],
            ["digital_will_items",
             "id (uuid), UNIQUE (user_id, scan_result_id)",
             "action, transfer_email, note"],
        ]
    )

    add_heading(doc, "14.3  API-Routen", level=2, color=BRAND_DARK)
    add_feature_table(doc,
        headers=["Route", "Methode", "Funktion"],
        rows=[
            ["/api/imap-scan",       "POST",  "IMAP-Verbindung, Scan, gibt FoundService[] zurück"],
            ["/api/hibp",            "GET",   "Have I Been Pwned Breach-Check für eine Domain"],
            ["/api/auth/callback",   "GET",   "Supabase Auth Callback (OAuth-Flow)"],
        ]
    )

    add_heading(doc, "14.4  Migrations-Übersicht", level=2, color=BRAND_DARK)
    add_feature_table(doc,
        headers=["Migration", "Inhalt"],
        rows=[
            ["20260506000000_initial_schema.sql",
             "profiles, scan_results, breach_alerts, RLS-Policies, handle_new_user Trigger"],
            ["20260507113000_add_scan_evidence_fields.sql",
             "evidence_count, evidence_types, sender_domains, detection_confidence, detection_source"],
            ["20260507120000_add_demo_detection_source.sql",
             "Erweiterung detection_source CHECK um 'demo' Wert"],
            ["20260507130000_digital_will.sql",
             "digital_will und digital_will_items Tabellen mit RLS"],
        ]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 15. PREISMODELLE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "15  Preismodelle – Free vs. Pro", level=1, color=BRAND_INDIGO)

    add_feature_table(doc,
        headers=["Feature", "Free", "Pro"],
        rows=[
            ["Account-Erkennung",            "✅ Ja",       "✅ Ja"],
            ["Max. angezeigte Accounts",      "20 Accounts", "Unbegrenzt"],
            ["Risiko-Score",                  "✅ Ja",       "✅ Ja"],
            ["Digital Twin (Netzwerk)",       "✅ Ja",       "✅ Ja"],
            ["Digital Will",                  "✅ Ja",       "✅ Ja"],
            ["Risiko-Prognose",               "✅ Ja",       "✅ Ja"],
            ["JSON/CSV-Export",               "✅ Ja",       "✅ Ja"],
            ["Einfache Löschlinks",           "✅ Ja",       "✅ Ja"],
            ["Automatisches Rescan",          "❌ Nein",     "✅ Ja"],
            ["Aktives Breach-Monitoring",     "❌ Nein",     "✅ Ja"],
            ["Breach-Alert-Benachrichtigungen","❌ Nein",     "✅ Ja"],
            ["Unbegrenzte Account-Anzeige",   "❌ Nein",     "✅ Ja"],
            ["Preis",                         "0 € / Monat", "Pro (Preis auf Anfrage)"],
        ]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 16. GLOSSAR
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "16  Glossar", level=1, color=BRAND_INDIGO)

    glossary = [
        ("Breach / Datenleck",
         "Ein Sicherheitsvorfall, bei dem Nutzerdaten eines Online-Dienstes gestohlen "
         "oder öffentlich zugänglich gemacht wurden."),
        ("DSGVO",
         "Datenschutz-Grundverordnung – EU-Datenschutzgesetz, das Rechte auf Datenportabilität, "
         "Löschung und Auskunft garantiert."),
        ("Digital Will",
         "GhostAccounts-Feature zur Regelung des digitalen Nachlasses bei Inaktivität oder Tod."),
        ("Evidence / Evidenz",
         "E-Mail-Metadaten, die beweisen, dass ein Account bei einem Dienst existiert "
         "(Registrierungs-, Sicherheits- oder Rechnungs-E-Mails)."),
        ("HIBP",
         "Have I Been Pwned – Dienst von Troy Hunt zum Prüfen ob eine E-Mail in Datenlecks vorkam."),
        ("IMAP",
         "Internet Message Access Protocol – Standardprotokoll für den Zugriff auf E-Mail-Postfächer."),
        ("OAuth 2.0 + PKCE",
         "Sicheres Authentifizierungsprotokoll ohne Passwort-Übertragung. "
         "PKCE verhindert Authorization Code Interception."),
        ("RLS",
         "Row Level Security – Datenbankfunktion, die sicherstellt dass Nutzer nur ihre eigenen Daten sehen."),
        ("Risiko-Score",
         "Numerischer Wert 0–100, der den Sicherheitsstatus des digitalen Fußabdrucks beschreibt."),
        ("Scan",
         "Analyse von E-Mail-Metadaten zur Erkennung von Online-Accounts."),
        ("Service-Domain",
         "Die Domain des erkannten Online-Dienstes (z.B. 'amazon.de', 'netflix.com')."),
        ("Digital Twin",
         "Interaktive D3.js Netzwerk-Visualisierung aller erkannten Accounts."),
    ]

    tbl = doc.add_table(rows=len(glossary), cols=2)
    tbl.style = "Table Grid"
    for ri, (term, definition) in enumerate(glossary):
        cells = tbl.rows[ri].cells
        bg = "EEF2FF" if ri % 2 == 0 else "FFFFFF"
        set_cell_bg(cells[0], bg)
        set_cell_bg(cells[1], bg)
        cells[0].width = Inches(1.8)
        cells[1].width = Inches(4.5)

        r0 = cells[0].paragraphs[0].add_run(term)
        r0.font.name = "Calibri"; r0.font.size = Pt(10); r0.font.bold = True
        r0.font.color.rgb = BRAND_INDIGO

        r1 = cells[1].paragraphs[0].add_run(definition)
        r1.font.name = "Calibri"; r1.font.size = Pt(10)
        r1.font.color.rgb = TEXT_DARK

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # 17. RECHTLICHE HINWEISE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "17  Rechtliche Hinweise", level=1, color=BRAND_INDIGO)

    add_heading(doc, "17.1  Betreiber", level=2, color=BRAND_DARK)
    add_info_box(doc, "WAMOCON GmbH", [
        "Mergenthalerallee 79–81",
        "65760 Eschborn, Deutschland",
        "Tel: +49 6196 5838311",
        "E-Mail: info@wamocon.com",
        "HRB 123666 · USt-ID: DE344930486",
    ], bg_hex="EEF2FF", border_hex="6366F1", title_color=BRAND_INDIGO)

    add_heading(doc, "17.2  Haftungshinweis", level=2, color=BRAND_DARK)
    add_body(doc,
        "GhostAccounts stellt Informationen über erkannte Online-Accounts bereit. "
        "Die Erkennung basiert auf E-Mail-Metadaten und kann in Einzelfällen unvollständig "
        "oder fehlerhaft sein. WAMOCON GmbH übernimmt keine Haftung für Entscheidungen, "
        "die auf Basis der Scan-Ergebnisse getroffen werden.")

    add_heading(doc, "17.3  Nutzungshinweise", level=2, color=BRAND_DARK)
    add_bullet(doc,
               "GhostAccounts darf ausschließlich für eigene E-Mail-Postfächer verwendet werden.")
    add_bullet(doc,
               "Das Scannen fremder Postfächer ohne ausdrückliche Genehmigung ist untersagt "
               "und kann strafbar sein.")
    add_bullet(doc,
               "Die Verwendung von App-Passwörtern ist bei entsprechenden Anbietern "
               "aus Sicherheitsgründen obligatorisch und in der Verantwortung des Nutzers.")

    add_heading(doc, "17.4  Rechtsdokumente", level=2, color=BRAND_DARK)
    add_body(doc, "Alle Rechtsdokumente sind im Footer der Anwendung verlinkt:")
    add_bullet(doc, "Impressum:  /de/legal/impressum")
    add_bullet(doc, "Datenschutzerklärung:  /de/legal/datenschutz")
    add_bullet(doc, "Allgemeine Geschäftsbedingungen (AGB):  /de/legal/agb")

    doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end_p.add_run(
        f"─── Ende des Produkthandbuchs ───\n"
        f"GhostAccounts v1.0  ·  WAMOCON GmbH  ·  {datetime.today().strftime('%d. %B %Y')}"
    )
    er.font.name = "Calibri"
    er.font.size = Pt(10)
    er.font.color.rgb = TEXT_MUTED
    er.font.italic = True

    return doc


# ── Ausführung ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    OUT_DIR = os.path.join(BASE, "Dokumente")
    os.makedirs(OUT_DIR, exist_ok=True)

    OUT_FILE = os.path.join(OUT_DIR, "GhostAccounts_Produkthandbuch.docx")

    print("🔨  Erstelle Produkthandbuch …")
    document = build_document()
    document.save(OUT_FILE)
    size_kb = os.path.getsize(OUT_FILE) / 1024
    print(f"✅  Gespeichert: {OUT_FILE}")
    print(f"📄  Dateigröße:  {size_kb:.1f} KB")
    print(f"🗂️   Ordner:      {OUT_DIR}")
